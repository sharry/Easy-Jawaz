import tkinter as tk
from tkinter import filedialog, ttk
from tkinter import *
from tkinter.ttk import *
import os
from pyzbar.pyzbar import decode
import cv2
from docx import Document
from docx.shared import Inches, Mm
from PIL import Image, ImageOps
import win32api, win32print
import qrcode
import qrcode.image.svg
from svglib.svglib import svg2rlg
from reportlab.graphics import renderPM
from pdf2image import convert_from_path
import subprocess
import psutil
import json
import windnd

printers = []
proc = subprocess.Popen(["wmic", "printer", "get", "name"], stdout=subprocess.PIPE, shell=True)
(out, err) = proc.communicate()
dataprinters = out.decode('utf-8').split('\r\r\n')
dataprinters = dataprinters[1:-1]
printers.clear()
appath = __file__[:-6]
for printer in dataprinters:
    if printer.strip() != "":
        printers.append(printer.strip())

root = tk.Tk()
root.resizable(False, False)
root.title('Easy Jawaz')
root.iconbitmap(appath + 'logo.ico')
jawazat = []
dataf = open(appath + 'settings.json',)
settings = json.load(dataf)
landscape = Image.open(appath + "landscape.jpg")
landscapex = Image.open(appath + "landscape_ex.jpg")
photophile = Image.open(appath + "photophile.jpg")

def draggedJawazat(files):
    for widget in frame.winfo_children():
        widget.destroy()
    label = tk.Label(frame, text="L'jawazat lli khtariti:", bg="#2D2D30", fg="white")
    label.pack()
    if jawazat != "":
        for file in files:
            im = Image.open(file)
            width, height = im.size
            if width == 450 and (height == 747 or height == 768):
                jawazat.append(file.decode("utf-8"))
            else:
                tk.messagebox.showinfo(title="Tswèra ma mqboulach!", message="Tswèra lli 3zelti dyal l'jawaz machi qedd (450x747) wella (450x768), ya imma qaddha wella hezzha kifma téléchager-tiha menn 'site")
    for jawaz in jawazat:
        label = tk.Label(frame, text=jawaz, bg="#007ACC", fg="white", cursor="hand2")
        label.bind("<Button-1>", OpenImage)
        label.pack(pady=2)



windnd.hook_dropfiles(root, draggedJawazat)
def isRunning(processName):
    for proc in psutil.process_iter():
        try:
            if processName.lower() in proc.name().lower():
                return True
        except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
            pass
    return False

def Clear():
    jawazat.clear()
    for widget in frame.winfo_children():
        widget.destroy()
    label = tk.Label(frame, text="L'jawazat lli khtariti:", bg="#2D2D30", fg="white")
    label.pack()

def OpenFile():
    files = filedialog.askopenfilenames(initialdir=os.path.join( os.getenv('USERPROFILE'), 'Downloads'),
    title="3zel l'jawaz lli ghat'imprimer",
    filetypes=(("images", "*.jpg"), ("all files", "*.*")))
    for widget in frame.winfo_children():
        widget.destroy()
    label = tk.Label(frame, text="L'jawazat lli khtariti:", bg="#2D2D30", fg="white")
    label.pack()
    if files != "":
        for file in files:
            im = Image.open(file)
            width, height = im.size
            if width == 450 and (height == 747 or height == 768):
                jawazat.append(file)
            else:
                tk.messagebox.showinfo(title="Tswèra ma mqboulach!", message="Tswèra lli 3zelti dyal l'jawaz machi qedd 450x747, ya imma qaddha wella hezzha kifma téléchager-tiha menn 'site")
    for jawaz in jawazat:
        label = tk.Label(frame, text=jawaz, bg="#007ACC", fg="white", cursor="hand2")
        label.bind("<Button-1>", OpenImage)
        label.pack(pady=2)

def OpenImage(event):
    os.startfile(event.widget['text'])

def OpenReady():
    os.startfile(appath + "ready.docx")

def uniquify(path):
    filename, extension = os.path.splitext(path)
    counter = 1

    while os.path.exists(path):
        path = filename + " (" + str(counter) + ")" + extension
        counter += 1

    return path

def ExportLandscape():
    if len(jawazat) > 0:
        for jawaz in jawazat:
            if ai.get() == 1:
                tswera = filedialog.askopenfilename(initialdir=os.path.join(os.getenv('USERPROFILE'), 'Desktop'),
                title="3zel 'tswèra",
                filetypes=(("images", "*.jpg"), ("images", "*.png"), ("all files", "*.*")))
                photo = Image.open(tswera)
                origin = Image.open(jawaz)
                phile = Image.open(appath + "photophile.jpg")
                prefilled = phile.copy()
                photo = photo.resize((170, 235))
                prefilled.paste(photo, (32, 0))
                postfilled = origin.copy()
                postfilled.paste(prefilled, (108, 132))
                postfilled = postfilled.convert('RGB')
                postfilled.save(appath + 'temp.jpg')
                original = Image.open(appath + 'temp.jpg')
            else:
                original = Image.open(jawaz)
            try:
                im = cv2.imread(jawaz)
                d = decode(im)
                data = d[0].data.decode('ascii')
                img = qrcode.make(data, image_factory = qrcode.image.svg.SvgImage)
                img.save(appath + "qrcode.svg")
                drawing = svg2rlg(appath + "qrcode.svg")
                renderPM.drawToFile(drawing, appath + "file.png", fmt="PNG")
            except:
                reply = tk.messagebox.askyesno(title="L'QR code mabghach yt'scana!", message="Had l'QR code lli fel'jawaz mabghach yt'scana.\nIla 3endek l'PDF dyalou dir 'Ah' w 3ezlou bach tjbed l'QR code mennou.\nIla ma3endekch dir 'Lla' w tqder tsawbou menn be3d bla QR code.")
                if reply == True:
                    pdfile = filedialog.askopenfilename(initialdir=os.path.join( os.getenv('USERPROFILE'), 'Downloads'),
                    title="3zel l'PDF lli yjbed mennou l'QR code",
                    filetypes=(("PDFs", "*.pdf"), ("all files", "*.*")))
                    pdf = convert_from_path(pdfile, 500, poppler_path = appath + 'poppler\\bin')
                    for page in pdf:
                        page.save(appath + 'hqr.jpg', 'JPEG')
                    cropping = Image.open(appath + 'hqr.jpg')
                    cropping = cropping.crop((300, 3720, 1400, 4820))
                    cropping.save(appath + 'hqr.jpg')
                    imx = cv2.imread(appath + 'hqr.jpg')
                    d = decode(imx)
                    data = d[0].data.decode('ascii')
                    tmpimg = qrcode.make(data, image_factory = qrcode.image.svg.SvgImage)
                    tmpimg.save(appath + "qrcode.svg")
                    drawing = svg2rlg(appath + "qrcode.svg")
                    renderPM.drawToFile(drawing, appath + "file.png", fmt="PNG")
                    os.remove(appath + "hqr.jpg")
                else:
                    return
            try:
                os.mkdir(os.getenv('USERPROFILE') + '\\Desktop\\' + os.path.basename(jawaz)[:-4])
            except:
                pass
            height = original.height
            if height == 747:
                qrcodd = original.crop((120, 145, 330, 355))
                person = original.crop((15, 400, 435, 530))
                vaccin = original.crop((15, 570, 435, 640))
                barcode = original.crop((170, 657, 335, 707))
                filled = landscape.copy()
                filled.paste(qrcodd, (30, 130))
                filled.paste(person, (290, 130))
                filled.paste(vaccin, (290, 280))
                filled.paste(barcode, (320, 385))
                filled = filled.convert('RGB')
                filled.save(uniquify(os.getenv('USERPROFILE') + '\\Desktop\\' + os.path.basename(jawaz)[:-4] + '\\Paysage.jpg'), quality=100)
                back = Image.open(appath + "l_exportable.jpg")
            elif height == 768:
                qrcodd = original.crop((120, 140, 330, 365))
                person = original.crop((15, 400, 435, 560))
                vaccin = original.crop((15, 580, 435, 660))
                barcode = original.crop((170, 680, 335, 727))
                filled = landscapex.copy()
                filled.paste(qrcodd, (30, 130))
                filled.paste(person, (290, 130))
                filled.paste(vaccin, (290, 280))
                filled.paste(barcode, (320, 400))
                filled = filled.convert('RGB')
                filled.save(uniquify(os.getenv('USERPROFILE') + '\\Desktop\\' + os.path.basename(jawaz)[:-4] + '\\Paysage.jpg'), quality=100)
                back = Image.open(appath + "l_exportable_ex.jpg")
            qrcd = Image.open(appath + "file.png")
            back_im = back.copy()
            back_im.paste(qrcd, (30, 50))
            back_im.save(uniquify(os.getenv('USERPROFILE') + '\\Desktop\\' + os.path.basename(jawaz)[:-4] + '\\Paysage_Back.jpg'), quality=100)
            os.remove(appath + "file.png")
            os.remove(appath + "qrcode.svg")
            try:
                os.remove(appath + "temp.jpg")
            except:
                pass
            os.startfile(os.path.realpath(os.getenv('USERPROFILE') + '\\Desktop\\' + os.path.basename(jawaz)[:-4]))
        Clear()
    else:
        tk.messagebox.showinfo(title="Ta Jawaz ma m3zoul!", message="3afak 3zel Jawaz wa7ed 3ell l'aqall")

def ExportPortrait():
    if len(jawazat) > 0:
        for jawaz in jawazat:
            if ai.get() == 1:
                tswera = filedialog.askopenfilename(initialdir=os.path.join(os.getenv('USERPROFILE'), 'Desktop'),
                title="3zel 'tswèra",
                filetypes=(("images", "*.jpg"), ("images", "*.png"), ("all files", "*.*")))
                photo = Image.open(tswera)
                origin = Image.open(jawaz)
                phile = Image.open(appath + "photophile.jpg")
                prefilled = phile.copy()
                photo = photo.resize((170, 235))
                prefilled.paste(photo, (32, 0))
                postfilled = origin.copy()
                postfilled.paste(prefilled, (108, 132))
                postfilled = postfilled.convert('RGB')
                postfilled.save(appath + 'temp.jpg')
                original = Image.open(appath + 'temp.jpg')
            else:
                original = Image.open(jawaz)
            try:
                im = cv2.imread(jawaz)
                d = decode(im)
                data = d[0].data.decode('ascii')
                img = qrcode.make(data, image_factory = qrcode.image.svg.SvgImage)
                img.save(appath + "qrcode.svg")
                drawing = svg2rlg(appath + "qrcode.svg")
                renderPM.drawToFile(drawing, appath + "file.png", fmt="PNG")
            except:
                reply = tk.messagebox.askyesno(title="L'QR code mabghach yt'scana!", message="Had l'QR code lli fel'jawaz mabghach yt'scana.\nIla 3endek l'PDF dyalou dir 'Ah' w 3ezlou bach tjbed l'QR code mennou.\nIla ma3endekch dir 'Lla' w tqder tsawbou menn be3d bla QR code.")
                if reply == True:
                    pdfile = filedialog.askopenfilename(initialdir=os.path.join( os.getenv('USERPROFILE'), 'Downloads'),
                    title="3zel l'PDF lli yjbed mennou l'QR code",
                    filetypes=(("PDFs", "*.pdf"), ("all files", "*.*")))
                    pdf = convert_from_path(pdfile, 500, poppler_path = appath + 'poppler\\bin')
                    for page in pdf:
                        page.save(appath + 'hqr.jpg', 'JPEG')
                    cropping = Image.open(appath + 'hqr.jpg')
                    cropping = cropping.crop((300, 3720, 1400, 4820))
                    cropping.save(appath + 'hqr.jpg')
                    imx = cv2.imread(appath + 'hqr.jpg')
                    d = decode(imx)
                    data = d[0].data.decode('ascii')
                    tmpimg = qrcode.make(data, image_factory = qrcode.image.svg.SvgImage)
                    tmpimg.save(appath + "qrcode.svg")
                    drawing = svg2rlg(appath + "qrcode.svg")
                    renderPM.drawToFile(drawing, appath + "file.png", fmt="PNG")
                    os.remove(appath + "hqr.jpg")
                else:
                    return
            try:
                os.mkdir(os.getenv('USERPROFILE') + '\\Desktop\\' + os.path.basename(jawaz)[:-4])
            except:
                pass
            original = original.convert('RGB')
            original.save(uniquify(os.getenv('USERPROFILE') + '\\Desktop\\' + os.path.basename(jawaz)[:-4] + '\\Paysage.jpg'), quality=100)
            height = original.height
            if height == 747:
                back = Image.open(appath + "p_exportable.jpg")
            elif height == 768:
                back = Image.open(appath + "p_exportable_ex.jpg")
            qrcd = Image.open(appath + "file.png")
            back_im = back.copy()
            back_im.paste(qrcd, (49, 29))
            back_im.save(uniquify(os.getenv('USERPROFILE') + '\\Desktop\\' + os.path.basename(jawaz)[:-4] + '\\Paysage_Back.jpg'), quality=100)
            os.remove(appath + "file.png")
            os.remove(appath + "qrcode.svg")
            try:
                os.remove(appath + "temp.jpg")
            except:
                pass
            os.startfile(os.path.realpath(os.getenv('USERPROFILE') + '\\Desktop\\' + os.path.basename(jawaz)[:-4]))
        Clear()
    else:
        tk.messagebox.showinfo(title="Ta Jawaz ma m3zoul!", message="3afak 3zel Jawaz wa7ed 3ell l'aqall")

def PrintWithQR():
    if (isRunning('WORD.EXE')):
        tk.messagebox.showinfo(title="Chi Word m7loul!", message="3afak sedd ga3 l'Word-at lli m7loulin, bach tqder t'imprimer.")
        return
    if len(jawazat) > 0:
        document = Document()
        section = document.sections[0]
        if psize.get() == 1:
            section.page_width = Mm(148)
            section.page_height = Mm(210)
            section.left_margin = Mm(12.7)
            section.right_margin = Mm(12.7)
            section.top_margin = Mm(12.7)
            section.bottom_margin = Mm(12.7)
            section.header_distance = Mm(12.7)
            section.footer_distance = Mm(12.7)
        else:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.left_margin = Mm(12.7)
            section.right_margin = Mm(12.7)
            section.top_margin = Mm(9)
            section.bottom_margin = Mm(12.7)
            section.header_distance = Mm(12.7)
            section.footer_distance = Mm(12.7)
        for jawaz in jawazat:
            if ai.get() == 1:
                tswera = filedialog.askopenfilename(initialdir=os.path.join(os.getenv('USERPROFILE'), 'Desktop'),
                title="3zel 'tswèra",
                filetypes=(("images", "*.jpg"), ("images", "*.png"), ("all files", "*.*")))
                photo = Image.open(tswera)
                original = Image.open(jawaz)
                phile = Image.open(appath + "photophile.jpg")
                prefilled = phile.copy()
                photo = photo.resize((170, 235))
                prefilled.paste(photo, (32, 0))
                filled = original.copy()
                filled.paste(prefilled, (108, 132))
                filled = filled.convert('RGB')
                img = filled
            else:
                img = Image.open(jawaz)
                img = img.convert('RGB')
            color = "#007ACC"
            border = (1, 1, 1, 1)
            im = cv2.imread(jawaz)
            try:
                d = decode(im)
                data = d[0].data.decode('ascii')
                tmpimg = qrcode.make(data, image_factory = qrcode.image.svg.SvgImage)
                tmpimg.save(appath + "qrcode.svg")
                drawing = svg2rlg(appath + "qrcode.svg")
                renderPM.drawToFile(drawing, appath + "file.png", fmt="PNG")
            except:
                reply = tk.messagebox.askyesno(title="L'QR code mabghach yt'scana!", message="Had l'QR code lli fel'jawaz mabghach yt'scana.\nIla 3endek l'PDF dyalou dir 'Ah' w 3ezlou bach tjbed l'QR code mennou.\nIla ma3endekch dir 'Lla' w tqder tsawbou menn be3d bla QR code.")
                if reply == True:
                    pdfile = filedialog.askopenfilename(initialdir=os.path.join( os.getenv('USERPROFILE'), 'Downloads'),
                    title="3zel l'PDF lli yjbed mennou l'QR code",
                    filetypes=(("PDFs", "*.pdf"), ("all files", "*.*")))
                    pdf = convert_from_path(pdfile, 500, poppler_path = appath + 'poppler\\bin')
                    for page in pdf:
                        page.save(appath + 'hqr.jpg', 'JPEG')
                    cropping = Image.open(appath + 'hqr.jpg')
                    cropping = cropping.crop((300, 3720, 1400, 4820))
                    cropping.save(appath + 'hqr.jpg')
                    imx = cv2.imread(appath + 'hqr.jpg')
                    d = decode(imx)
                    data = d[0].data.decode('ascii')
                    tmpimg = qrcode.make(data, image_factory = qrcode.image.svg.SvgImage)
                    tmpimg.save(appath + "qrcode.svg")
                    drawing = svg2rlg(appath + "qrcode.svg")
                    renderPM.drawToFile(drawing, appath + "file.png", fmt="PNG")
                    os.remove(appath + "hqr.jpg")
                else:
                    return
            height = img.height
            if v.get() == 1:
                if height == 747:
                    back = Image.open(appath + "p_back.jpg")
                    qrcd = Image.open(appath + "file.png")
                    back_im = back.copy()
                    back_im.paste(qrcd, (49, 30))
                    back_im.save(appath + "p_back.jpg", quality=100)
                elif height == 768:
                    back = Image.open(appath + "p_back_ex.jpg")
                    qrcd = Image.open(appath + "file.png")
                    back_im = back.copy()
                    back_im.paste(qrcd, (49, 30))
                    back_im.save(appath + "p_back_ex.jpg", quality=100)
                os.remove(appath + "file.png")
                os.remove(appath + "qrcode.svg")
                new_img = ImageOps.expand(img, border=border, fill=color)
                newf = jawaz.replace('.jpg', 'x.jpg')
                new_img.save(newf)
                im1 = cv2.imread(newf)
                if height == 747:
                    im2 = cv2.imread(appath + "p_back.jpg")
                elif height == 768:
                    im2 = cv2.imread(appath + "p_back_ex.jpg")
                im_h = cv2.hconcat([im1, im2])
                cv2.imwrite(newf, im_h)
            if v.get() == 2:
                if height == 747:
                    back = Image.open(appath + "l_back.jpg")
                    qrcd = Image.open(appath + "file.png")
                    qrcd = qrcd.rotate(90)
                    back_im = back.copy()
                    back_im.paste(qrcd, (49, 370))
                    back_im.save(appath + "l_back.jpg", quality=100)
                elif height == 768:
                    back = Image.open(appath + "l_back_ex.jpg")
                    qrcd = Image.open(appath + "file.png")
                    qrcd = qrcd.rotate(90)
                    back_im = back.copy()
                    back_im.paste(qrcd, (49, 370))
                    back_im.save(appath + "l_back_ex.jpg", quality=100)
                os.remove(appath + "file.png")
                os.remove(appath + "qrcode.svg")
                if height == 747:
                    qrcodd = img.crop((120, 145, 330, 355))
                    person = img.crop((15, 400, 435, 530))
                    vaccin = img.crop((15, 570, 435, 640))
                    barcode = img.crop((170, 657, 335, 707))
                    filled = landscape.copy()
                    filled.paste(qrcodd, (30, 130))
                    filled.paste(person, (290, 130))
                    filled.paste(vaccin, (290, 280))
                    filled.paste(barcode, (320, 385))
                elif height == 768:
                    qrcodd = img.crop((120, 140, 330, 365))
                    person = img.crop((15, 400, 435, 560))
                    vaccin = img.crop((15, 580, 435, 660))
                    barcode = img.crop((170, 680, 335, 727))
                    filled = landscapex.copy()
                    filled.paste(qrcodd, (30, 130))
                    filled.paste(person, (290, 130))
                    filled.paste(vaccin, (290, 280))
                    filled.paste(barcode, (320, 400))
                filled = filled.convert('RGB')
                new_img = ImageOps.expand(filled, border=border, fill=color)
                newf = jawaz.replace('.jpg', 'x.jpg')
                new_img.save(newf, quality=100)
                im = Image.open(newf)
                out = im.rotate(-90, expand=True)
                out.save(newf)
                im1 = cv2.imread(newf)
                if height == 747:
                    im2 = cv2.imread(appath + "l_back.jpg")
                elif height == 768:
                    im2 = cv2.imread(appath + "l_back_ex.jpg")
                im_h = cv2.hconcat([im1, im2])
                cv2.imwrite(newf, im_h)
            document.add_picture(newf, width=Inches(4))
            os.remove(newf)
        document.save(appath + "ready.docx")
        Clear()
        if dp.get() == 1:
            os.startfile(appath + "ready.docx")
        else:
            os.system('cmd /c "wmic printer where name="' + pl.get() + '" call setdefaultprinter"')
            win32api.ShellExecute (0,
            "print",
            appath + 'ready.docx',
            '/d:"%s"' % win32print.GetDefaultPrinter (),
            ".",
            0)
    else:
        tk.messagebox.showinfo(title="Ta Jawaz ma m3zoul!", message="3afak 3zel Jawaz wa7ed 3ell l'aqall")

def PrintWithout():
    if (isRunning('WORD.EXE')):
        tk.messagebox.showinfo(title="Chi Word m7loul!", message="3afak sedd ga3 l'Word-at lli m7loulin, bach tqder t'imprimer.")
        return
    if len(jawazat) > 0:
        document = Document()
        section = document.sections[0]
        if psize.get() == 1:
            section.page_width = Mm(148)
            section.page_height = Mm(210)
            section.left_margin = Mm(12.7)
            section.right_margin = Mm(12.7)
            section.top_margin = Mm(12.7)
            section.bottom_margin = Mm(12.7)
            section.header_distance = Mm(12.7)
            section.footer_distance = Mm(12.7)
        else:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.left_margin = Mm(12.7)
            section.right_margin = Mm(12.7)
            section.top_margin = Mm(9)
            section.bottom_margin = Mm(12.7)
            section.header_distance = Mm(12.7)
            section.footer_distance = Mm(12.7)
        for jawaz in jawazat:
            if ai.get() == 1:
                tswera = filedialog.askopenfilename(initialdir=os.path.join(os.getenv('USERPROFILE'), 'Desktop'),
                title="3zel 'tswèra",
                filetypes=(("images", "*.jpg"), ("images", "*.png"), ("all files", "*.*")))
                photo = Image.open(tswera)
                original = Image.open(jawaz)
                phile = Image.open(appath + "photophile.jpg")
                prefilled = phile.copy()
                photo = photo.resize((170, 235))
                prefilled.paste(photo, (32, 0))
                filled = original.copy()
                filled.paste(prefilled, (108, 132))
                filled = filled.convert('RGB')
                img = filled
            else:
                img = Image.open(jawaz)
                img = img.convert('RGB')
            color = "#007ACC"
            border = (1, 1, 1, 1)
            height = img.height
            if v.get() == 1:
                new_img = ImageOps.expand(img, border=border, fill=color)
                newf = jawaz.replace('.jpg', 'x.jpg')
                new_img.save(newf)
                im1 = cv2.imread(newf)
                if height == 747:
                    im2 = cv2.imread(appath + "p_without.jpg")
                elif height == 768:
                    im2 = cv2.imread(appath + "p_without_ex.jpg")
                im_h = cv2.hconcat([im1, im2])
                cv2.imwrite(newf, im_h)
            if v.get() == 2:
                if height == 747:
                    qrcodd = img.crop((120, 145, 330, 355))
                    person = img.crop((15, 400, 435, 530))
                    vaccin = img.crop((15, 570, 435, 640))
                    barcode = img.crop((170, 657, 335, 707))
                    filled = landscape.copy()
                    filled.paste(qrcodd, (30, 130))
                    filled.paste(person, (290, 130))
                    filled.paste(vaccin, (290, 280))
                    filled.paste(barcode, (320, 385))
                elif height == 768:
                    qrcodd = img.crop((120, 140, 330, 365))
                    person = img.crop((15, 400, 435, 560))
                    vaccin = img.crop((15, 580, 435, 660))
                    barcode = img.crop((170, 680, 335, 727))
                    filled = landscapex.copy()
                    filled.paste(qrcodd, (30, 130))
                    filled.paste(person, (290, 130))
                    filled.paste(vaccin, (290, 280))
                    filled.paste(barcode, (320, 400))
                filled = filled.convert('RGB')
                new_img = ImageOps.expand(filled, border=border, fill=color)
                newf = jawaz.replace('.jpg', 'x.jpg')
                new_img.save(newf, quality=100)
                im = Image.open(newf)
                out = im.rotate(-90, expand=True)
                out.save(newf)
                im1 = cv2.imread(newf)
                if height == 747:
                    im2 = cv2.imread(appath + "l_without.jpg")
                if height == 768:
                    im2 = cv2.imread(appath + "l_without_ex.jpg")
                im_h = cv2.hconcat([im1, im2])
                cv2.imwrite(newf, im_h)
            document.add_picture(newf, width=Inches(4))
            os.remove(newf)
        document.save(appath + "ready.docx")
        Clear()
        if dp.get() == 1:
            os.startfile(appath + "ready.docx")
        else:
            os.system('cmd /c "wmic printer where name="' + pl.get() + '" call setdefaultprinter"')
            win32api.ShellExecute (0,
            "print",
            appath + 'ready.docx',
            '/d:"%s"' % win32print.GetDefaultPrinter (),
            ".",
            0)
    else:
        tk.messagebox.showinfo(title="Ta Jawaz ma m3zoul!", message="3afak 3zel Jawaz wa7ed 3ell l'aqall")

def Kebber():
    if len(jawazat) > 0:
        for jawaz in jawazat:
            subject = Image.open(jawaz)
            height = subject.height
            result = subject.copy()
            if height == 768:
                cine = subject.crop((190, 400, 265, 420))
                cine = cine.resize((113, 30))
                name = subject.crop((140, 445, 310, 505))
                name = name.resize((250, 88))
                birt = subject.crop((185, 525, 265, 545))
                birt = birt.resize((120, 30))
                date = subject.crop((185, 595, 265, 610))
                date = date.resize((120, 23))
                result.paste(cine, (180, 395))
                result.paste(name, (105, 425))
                result.paste(birt, (180, 525))
                result.paste(date, (180, 592))
            elif height == 747:
                cine = subject.crop((195, 400, 280, 420))
                cine = cine.resize((128, 30))
                name = subject.crop((135, 445, 310, 485))
                name = name.resize((262, 60))
                birt = subject.crop((185, 505, 265, 525))
                birt = birt.resize((120, 30))
                date = subject.crop((185, 575, 265, 595))
                date = date.resize((120, 30))
                result.paste(cine, (180, 395))
                result.paste(name, (98, 430))
                result.paste(birt, (180, 505))
                result.paste(date, (180, 570))
            result = result.convert('RGB')
            result.save(jawaz)
        tk.messagebox.showinfo(title="Safi", message="Ra l'ktaba wellat daba kbira")
    else:
        tk.messagebox.showinfo(title="Ta Jawaz ma m3zoul!", message="3afak 3zel Jawaz wa7ed 3ell l'aqall")


def About():
    tk.messagebox.showinfo(title="Contact", message="\tSharry Automation © 2021                       \n\n\tPhone:\t(+212) 07 00 45 90 37\n\tEmail:\tsharryautomation@gmail.com\n\tPage:\tfb.com/SharryAutomation")
    return

def UpdateSettings():
    dataf = open(appath + 'settings.json', "w")
    for setting in settings['settings']:
        setting['Yzdts'] = ai.get()
        setting['Yimpr'] = dp.get()
        setting['Lqedd'] = psize.get()
    json.dump(settings, dataf)
    dataf.close()

def PrinterChanged(event):
    dataf = open(appath + 'settings.json', "w")
    for setting in settings['settings']:
        setting['Curnt'] = event.widget.get()
    json.dump(settings, dataf)
    dataf.close()

canvas = tk.Canvas(root, height=500, width=500)
canvas.pack()

frame = tk.Frame(root, bg="#2D2D30")
frame.place(relwidth=0.8, relheight=0.3, relx=0.1, rely=0.06)
label = tk.Label(frame, text="L'jawazat lli khtariti:", bg="#2D2D30", fg="white")
label.pack()
clear = tk.Button(canvas, text="7eyyed koullchi", fg="white", bg="#D61424", command=Clear)
clear.place(relx=0.515, rely=0.38)
openFile = tk.Button(canvas, text="   3zel l'jawaz   ", fg="white", bg="#007ACC", command=OpenFile)
openFile.place(relx=0.335, rely=0.38)
print07dh = tk.Button(canvas, text=" Saweb bla QR Code ", fg="white", bg="#68217A", command=PrintWithout)
print07dh.place(relx=0.267, rely=0.45)
print10dh = tk.Button(canvas, text="  Saweb b'QR Code  ", fg="white", bg="#68217A", command=PrintWithQR)
print10dh.place(relx=0.518, rely=0.45)
ai = tk.IntVar(value=0)
addimage= tk.Checkbutton(canvas, text="Yzid lel'jawazat tswèra", variable=ai, onvalue=1, offvalue=0, command=UpdateSettings)
addimage.place(relx=0.1, rely=0.63)
dp = tk.IntVar(value=0)
dontprint = tk.Checkbutton(canvas, text="Ysaweb l'Word w bla ma y'imprimer", variable=dp, onvalue=1, offvalue=0, command=UpdateSettings)
dontprint.place(relx=0.1, rely=0.58)
abtimg = PhotoImage(file = appath + "about.png")
about = tk.Button(canvas, image=abtimg, relief='flat', command=About)
about.place(relheight=0.05, relwidth=0.05, relx=0.94, rely=0.01)
wrdimg = PhotoImage(file = appath + "word.png")
word = tk.Button(canvas, image=wrdimg, command=OpenReady)
word.place(relheight=0.2, relwidth=0.2, relx=0.68, rely=0.6)
exportportrait = tk.Button(canvas, text="Sedderhoum waqfin (Bureau)", fg="white", bg="#007ACC", command=ExportPortrait)
exportportrait.place(relwidth=0.33,relx=0.1, rely=0.69)
exportlandscape = tk.Button(canvas, text="Sedderhoum na3sin  (Bureau)", fg="white", bg="#007ACC", command=ExportLandscape)
exportlandscape.place(relwidth=0.33, relx=0.1, rely=0.745)
kebber = tk.Button(canvas, text="Kebber\nl'ktaba", fg="white", bg="#68217A", command=Kebber)
kebber.place(relwidth=0.11, relheight=0.11, relx=0.445, rely=0.69)
pl = tk.StringVar()
prlabel = tk.Label(canvas, text="L'imprimante lli ghatkellef:")
prlabel.place(relwidth=0.35, relx=0.067, rely=0.83)
v = tk.IntVar()
radwa = tk.Radiobutton(canvas, text="Waqfin", padx = 20, variable=v, value=1)
radna = tk.Radiobutton(canvas, text="Na3sin", padx = 20, variable=v, value=2)
radwa.place(relx=0.3, rely=0.51)
radna.place(relx=0.5, rely=0.51)
radwa.select()
printerslist = ttk.Combobox(canvas, state="readonly", textvariable = pl)
printerslist['values'] = tuple(printers)
printerslist.bind("<<ComboboxSelected>>", PrinterChanged)
printerslist.place(relwidth=0.45, relx=0.43, rely=0.83)
lblsize = tk.Label(canvas, text="L'qedd dyal l'werqa:")
lblsize.place(relx=0.095, rely=0.88)
psize = tk.IntVar()
a5 = tk.Radiobutton(canvas, text="A5", padx = 20, variable=psize, value=1, command=UpdateSettings)
a4 = tk.Radiobutton(canvas, text="A4", padx = 20, variable=psize, value=2, command=UpdateSettings)
a5.place(relx=0.4, rely=0.88)
a4.place(relx=0.55, rely=0.88)

currentpos = -1
for setting in settings['settings']:
    ai.set(setting['Yzdts'])
    dp.set(setting['Yimpr'])
    psize.set(setting['Lqedd'])
    for printer in printers:
        currentpos += 1
        if printer == setting['Curnt']:
            break
printerslist.current(currentpos)
dataf.close()
root.eval('tk::PlaceWindow . center')
root.mainloop()